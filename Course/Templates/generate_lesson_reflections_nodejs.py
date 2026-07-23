from __future__ import annotations

import argparse
import re
import shutil
import sys
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.document import Document as DocumentObject
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError as exc:
    raise SystemExit(
        "This script requires python-docx. Install it with:\n"
        "    py -m pip install python-docx"
    ) from e


DOCPAC_FOLDER_RE = re.compile(r"^docpac_(\d{6})$", re.IGNORECASE)
LESSON_FOLDER_RE = re.compile(r"^Unit(\d+)Lesson(\d+)$", re.IGNORECASE)
STEP_HEADING_RE = re.compile(r"^Step\s+(\d+)\b", re.IGNORECASE)
END_REFLECTION_RE = re.compile(
    r"End[\s\-–—]*of[\s\-–—]*Lesson\s+Reflections?\b",
    re.IGNORECASE,
)
STEP_REFLECTION_RE = re.compile(
    r"^(?:Step\s+)?Reflections?\s*[:\-–—]?\s*(.*)$",
    re.IGNORECASE,
)
TRAILING_REFLECTION_RE = re.compile(r"\bReflections?\s*[:\-–—]?\s*$", re.IGNORECASE)
NUMBER_PREFIX_RE = re.compile(r"^\s*(?:\(?\d+\)?[.)]|[-•])\s*")

# These sections normally follow reflection content in the lesson format.
SECTION_BOUNDARY_RE = re.compile(
    r"^(?:"
    r"Deliverables?|File Requirements?|Submission Requirements?|"
    r"Common Pitfalls?|Related CTE Tasks?|Related Tasks?|"
    r"Assessment|Grading|Extensions?|Bonus|Resources?|Teacher Notes?"
    r")\b",
    re.IGNORECASE,
)


@dataclass
class ReflectionQuestions:
    step_questions: list[tuple[int, str]]
    end_questions: list[str]


@dataclass
class TemplateParts:
    title: object
    step_heading: object
    question: object
    answer_table: object
    end_heading: object


def parse_args() -> argparse.Namespace:
    script_dir = Path(__file__).resolve().parent

    parser = argparse.ArgumentParser(
        description=(
            "Traverse docpac_XXXXXX folders, extract lesson reflection questions, "
            "and create printable reflection documents from a Word template."
        )
    )
    parser.add_argument(
        "--root",
        type=Path,
        default=script_dir,
        help="Folder containing the docpac_XXXXXX folders (default: script folder).",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=script_dir / "Reflections Lessons.docx",
        help="Reflection worksheet template.",
    )
    parser.add_argument(
        "--output-pattern",
        default="{lesson}Reflections.docx",
        help=(
            "Output filename pattern. Available fields: {lesson}, {unit}, "
            "{lesson_number}, and {docpac}. "
            "Default: {lesson}Reflections.docx"
        ),
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show what would be created without writing files.",
    )
    return parser.parse_args()


def clean_text(text: str) -> str:
    """Normalize Word line breaks and repeated whitespace without changing wording."""
    text = text.replace("\u00ad", "").replace("\u200b", "")
    return re.sub(r"\s+", " ", text).strip()


def strip_number_prefix(text: str) -> str:
    return NUMBER_PREFIX_RE.sub("", text, count=1).strip()


def is_intro_instruction(text: str) -> bool:
    """Identify directions that introduce reflection questions but are not questions."""
    normalized = clean_text(text).casefold().rstrip(".:; ")

    instruction_starts = (
        "answer ",
        "respond ",
        "complete ",
        "write ",
        "use ",
        "address ",
        "consider ",
    )
    instruction_phrases = (
        "the following question",
        "the following questions",
        "these question",
        "these questions",
        "in the assignment submission",
        "in your assignment submission",
        "in your submission",
        "below",
    )

    return normalized.startswith(instruction_starts) and any(
        phrase in normalized for phrase in instruction_phrases
    )


def is_heading_paragraph(paragraph: Paragraph) -> bool:
    try:
        style_name = paragraph.style.name.casefold()
    except (AttributeError, KeyError):
        return False
    return style_name.startswith("heading") or style_name == "title"


def is_collection_boundary(paragraph: Paragraph, text: str) -> bool:
    if STEP_HEADING_RE.match(text):
        return True
    if END_REFLECTION_RE.search(text):
        return True
    if SECTION_BOUNDARY_RE.match(text):
        return True

    # The lesson files use heading styles for major sections. Once reflection
    # collection has begun, an unrelated heading marks the end of that block.
    if is_heading_paragraph(paragraph):
        return True

    return False


def add_question(target: list[str], text: str, *, strip_number: bool = False) -> None:
    text = clean_text(text)
    if strip_number:
        text = strip_number_prefix(text)
    if not text or is_intro_instruction(text):
        return
    target.append(text)


def extract_reflections(source_path: Path) -> ReflectionQuestions:
    document = Document(source_path)

    step_questions: list[tuple[int, str]] = []
    end_questions: list[str] = []

    current_step: int | None = None
    mode: str | None = None  # None, "step", or "end"

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue

        # Check this first because the phrase also contains the word Reflection.
        end_match = END_REFLECTION_RE.search(text)
        if end_match:
            mode = "end"
            current_step = None
            remainder = text[end_match.end() :].strip(" :\t-–—")
            add_question(end_questions, remainder, strip_number=True)
            continue

        step_match = STEP_HEADING_RE.match(text)
        if step_match:
            current_step = int(step_match.group(1))
            mode = None
            continue

        if mode == "end":
            if is_collection_boundary(paragraph, text):
                mode = None
                continue
            add_question(end_questions, text, strip_number=True)
            continue

        if mode == "step":
            if is_collection_boundary(paragraph, text):
                mode = None
                # A new step was already handled above. Other section headings
                # only terminate reflection collection.
                continue
            if current_step is not None:
                cleaned = strip_number_prefix(text)
                if cleaned and not is_intro_instruction(cleaned):
                    step_questions.append((current_step, cleaned))
            continue

        if current_step is None:
            continue

        # Normal case: a paragraph containing only "Reflection" or
        # "Reflections", optionally followed by a question on the same line.
        reflection_match = STEP_REFLECTION_RE.match(text)
        if reflection_match:
            mode = "step"
            remainder = reflection_match.group(1).strip()
            if remainder:
                step_questions.append((current_step, strip_number_prefix(remainder)))
            continue

        # Defensive case for files where the Reflection heading was accidentally
        # joined to the end of the preceding paragraph during editing.
        if TRAILING_REFLECTION_RE.search(text):
            mode = "step"

    return ReflectionQuestions(step_questions, end_questions)


def find_direct_paragraph_element(document: DocumentObject, text: str) -> object:
    wanted = text.casefold()
    for paragraph in document.paragraphs:
        if clean_text(paragraph.text).casefold() == wanted:
            return paragraph._p
    raise RuntimeError(f"Template paragraph not found: {text!r}")


def find_template_parts(document: DocumentObject) -> TemplateParts:
    body = document._element.body
    children = list(body)

    title_element = None
    step_heading_element = None
    end_heading_element = None

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if title_element is None and LESSON_FOLDER_RE.search(text.replace(" ", "")):
            title_element = paragraph._p
        elif text.casefold() == "step reflections":
            step_heading_element = paragraph._p
        elif text.casefold() in {
            "end-of-lesson reflections",
            "end-of-lesson reflection",
        }:
            end_heading_element = paragraph._p

    # The supplied template title contains spaces, so the folder regex does not
    # match it directly. Fall back to its known placeholder wording.
    if title_element is None:
        title_element = find_direct_paragraph_element(document, "Unit 0 Lesson 0")
    if step_heading_element is None:
        step_heading_element = find_direct_paragraph_element(document, "Step Reflections")
    if end_heading_element is None:
        try:
            end_heading_element = find_direct_paragraph_element(
                document, "End-of-Lesson Reflections"
            )
        except RuntimeError:
            end_heading_element = find_direct_paragraph_element(
                document, "End-of-Lesson Reflection"
            )

    step_index = children.index(step_heading_element)
    end_index = children.index(end_heading_element)

    question_element = None
    answer_table_element = None
    for child in children[step_index + 1 : end_index]:
        if child.tag == qn("w:p") and question_element is None:
            paragraph = Paragraph(child, document._body)
            if clean_text(paragraph.text):
                question_element = child
        elif child.tag == qn("w:tbl") and question_element is not None:
            answer_table_element = child
            break

    if question_element is None or answer_table_element is None:
        raise RuntimeError(
            "The template must contain a sample question followed by an answer table."
        )

    return TemplateParts(
        title=deepcopy(title_element),
        step_heading=deepcopy(step_heading_element),
        question=deepcopy(question_element),
        answer_table=deepcopy(answer_table_element),
        end_heading=deepcopy(end_heading_element),
    )


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace visible text while retaining the sample paragraph's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_keep_with_next(paragraph: Paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value
    paragraph.paragraph_format.keep_together = True


def prevent_row_split(row: object) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def protect_question_block(question: Paragraph, answer_table: Table) -> None:
    """
    Keep a question and its complete answer area on the same page.

    Word's keep-with-next chain moves the full block to the next page when it
    cannot fit, while cantSplit prevents an individual answer row from breaking.
    """
    set_keep_with_next(question, True)

    rows = answer_table.rows
    for row_index, row in enumerate(rows):
        prevent_row_split(row)
        keep_next = row_index < len(rows) - 1
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = keep_next


def insert_before_section(document: DocumentObject, element: object) -> None:
    body = document._element.body
    section_properties = body.sectPr
    if section_properties is None:
        body.append(element)
    else:
        body.insert(body.index(section_properties), element)


def append_paragraph_from_template(
    document: DocumentObject,
    prototype: object,
    text: str,
) -> Paragraph:
    element = deepcopy(prototype)
    insert_before_section(document, element)
    paragraph = Paragraph(element, document._body)
    set_paragraph_text(paragraph, text)
    return paragraph


def append_table_from_template(
    document: DocumentObject,
    prototype: object,
) -> Table:
    element = deepcopy(prototype)
    insert_before_section(document, element)
    return Table(element, document._body)


def clear_template_body(document: DocumentObject) -> None:
    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def build_reflection_document(
    template_path: Path,
    output_path: Path,
    unit: int,
    lesson: int,
    questions: ReflectionQuestions,
) -> None:
    # Copy first so custom styles, theme data, and other package parts are kept.
    shutil.copy2(template_path, output_path)
    document = Document(output_path)
    parts = find_template_parts(document)
    clear_template_body(document)

    title = append_paragraph_from_template(
        document,
        parts.title,
        f"Unit {unit} Lesson {lesson}",
    )
    set_keep_with_next(title, True)

    if questions.step_questions:
        heading = append_paragraph_from_template(
            document,
            parts.step_heading,
            "Step Reflections",
        )
        set_keep_with_next(heading, True)

        for _step_number, question_text in questions.step_questions:
            question = append_paragraph_from_template(
                document,
                parts.question,
                question_text,
            )
            answer_table = append_table_from_template(document, parts.answer_table)
            protect_question_block(question, answer_table)

    if questions.end_questions:
        heading = append_paragraph_from_template(
            document,
            parts.end_heading,
            "End-of-Lesson Reflections",
        )
        set_keep_with_next(heading, True)

        for question_text in questions.end_questions:
            question = append_paragraph_from_template(
                document,
                parts.question,
                question_text,
            )
            answer_table = append_table_from_template(document, parts.answer_table)
            protect_question_block(question, answer_table)

    document.save(output_path)


def iter_lesson_folders(root: Path) -> Iterable[tuple[str, Path, int, int]]:
    """Yield lesson folders below each docpac_XXXXXX directory."""
    found: list[tuple[str, Path, int, int]] = []

    for docpac_folder in root.iterdir():
        if not docpac_folder.is_dir():
            continue
        docpac_match = DOCPAC_FOLDER_RE.fullmatch(docpac_folder.name)
        if not docpac_match:
            continue

        # rglob permits an extra organizational folder without changing the
        # expected UnitXLessonY naming requirement.
        for lesson_folder in docpac_folder.rglob("Unit*Lesson*"):
            if not lesson_folder.is_dir():
                continue
            lesson_match = LESSON_FOLDER_RE.fullmatch(lesson_folder.name)
            if not lesson_match:
                continue
            found.append(
                (
                    docpac_match.group(1),
                    lesson_folder,
                    int(lesson_match.group(1)),
                    int(lesson_match.group(2)),
                )
            )

    yield from sorted(found, key=lambda item: (item[0], item[2], item[3], str(item[1])))


def make_output_name(
    pattern: str,
    *,
    lesson_name: str,
    unit: int,
    lesson_number: int,
    docpac: str,
) -> str:
    try:
        filename = pattern.format(
            lesson=lesson_name,
            unit=unit,
            lesson_number=lesson_number,
            docpac=docpac,
        )
    except KeyError as exc:
        raise ValueError(f"Unknown output-pattern field: {exc.args[0]}") from exc

    if Path(filename).name != filename or not filename.lower().endswith(".docx"):
        raise ValueError("The output pattern must produce a .docx filename, not a path.")
    return filename


def main() -> int:
    args = parse_args()
    root = args.root.resolve()
    template_path = args.template.resolve()

    if not root.is_dir():
        print(f"Root folder does not exist: {root}", file=sys.stderr)
        return 1
    if not template_path.is_file():
        print(f"Template does not exist: {template_path}", file=sys.stderr)
        return 1

    created = 0
    skipped = 0
    failed = 0
    lesson_folders = list(iter_lesson_folders(root))

    if not lesson_folders:
        print(f"No UnitXLessonY folders were found below docpac folders in {root}.")
        return 2

    for docpac_code, lesson_folder, unit, lesson_number in lesson_folders:
        lesson_name = lesson_folder.name
        source_path = lesson_folder / f"{lesson_name}.docx"

        if not source_path.is_file():
            print(f"Skipped {source_path.relative_to(root)}: source lesson not found.")
            skipped += 1
            continue

        try:
            questions = extract_reflections(source_path)
            if not questions.step_questions and not questions.end_questions:
                print(
                    f"Skipped {source_path.relative_to(root)}: "
                    "no reflection questions were found."
                )
                skipped += 1
                continue

            output_name = make_output_name(
                args.output_pattern,
                lesson_name=lesson_name,
                unit=unit,
                lesson_number=lesson_number,
                docpac=docpac_code,
            )
            output_path = lesson_folder / output_name

            if args.dry_run:
                print(
                    f"Would create {output_path.relative_to(root)} "
                    f"({len(questions.step_questions)} step, "
                    f"{len(questions.end_questions)} end-of-lesson questions)."
                )
                created += 1
                continue

            build_reflection_document(
                template_path,
                output_path,
                unit,
                lesson_number,
                questions,
            )
            print(
                f"Created {output_path.relative_to(root)} "
                f"({len(questions.step_questions)} step, "
                f"{len(questions.end_questions)} end-of-lesson questions)."
            )
            created += 1

        except Exception as exc:
            print(
                f"Failed {source_path.relative_to(root)}: {exc}",
                file=sys.stderr,
            )
            failed += 1

    print(
        f"\nFinished: {created} created, {skipped} skipped, {failed} failed."
    )
    return 1 if failed else (0 if created else 2)


if __name__ == "__main__":
    raise SystemExit(main())
